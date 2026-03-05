"""DOCX generation tool using python-docx with BeautifulSoup HTML parsing."""

from pathlib import Path
from typing import Annotated

from agent_framework import ai_function
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from pydantic import Field

# Color constants matching the CSS template
COLOR_DARK_NAVY = RGBColor(0x0D, 0x1B, 0x2A)  # #0d1b2a - headings/strong
COLOR_MUTED_BLUE = RGBColor(0x41, 0x5A, 0x77)  # #415a77 - contact/dates
COLOR_DARK_BLUE = RGBColor(0x1B, 0x26, 0x3B)  # #1b263b - profile/company/h3
COLOR_BORDER = RGBColor(0x77, 0x8D, 0xA9)  # #778da9 - borders
COLOR_BODY = RGBColor(0x1A, 0x1A, 0x1A)  # #1a1a1a - body text


def _add_bottom_border(paragraph, color_hex: str, size: int = 4):
    """Add a bottom border to a paragraph."""
    pPr = paragraph.paragraph_format.element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:color"), color_hex)
    bottom.set(qn("w:space"), "1")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_hyperlink(paragraph, url: str, text: str):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "415a77")
    rPr.append(color)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _get_element_classes(tag: Tag) -> set:
    """Get CSS classes of an element."""
    classes = tag.get("class", [])
    if isinstance(classes, str):
        classes = classes.split()
    return set(classes)


def _is_contact(tag: Tag) -> bool:
    """Check if element has contact-related class."""
    classes = _get_element_classes(tag)
    return bool(classes & {"contact", "contact-info"})


def _is_profile(tag: Tag) -> bool:
    """Check if element has profile/summary class."""
    classes = _get_element_classes(tag)
    return bool(classes & {"profile", "summary"})


def _add_inline_runs(paragraph, element, bold=False, italic=False, font_name=None, font_size=None, color=None):
    """Recursively add inline content from an element to a paragraph."""
    for child in element.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text:
                run = paragraph.add_run(text)
                run.bold = bold
                run.italic = italic
                if font_name:
                    run.font.name = font_name
                if font_size:
                    run.font.size = font_size
                if color:
                    run.font.color.rgb = color
        elif isinstance(child, Tag):
            if child.name == "br":
                run = paragraph.add_run()
                run.add_break()
            elif child.name in ("strong", "b"):
                _add_inline_runs(
                    paragraph, child,
                    bold=True, italic=italic,
                    font_name=font_name, font_size=font_size,
                    color=color or COLOR_DARK_NAVY,
                )
            elif child.name in ("em", "i"):
                _add_inline_runs(
                    paragraph, child,
                    bold=bold, italic=True,
                    font_name=font_name, font_size=font_size,
                    color=color,
                )
            elif child.name == "a":
                href = child.get("href", "")
                link_text = child.get_text()
                if href:
                    _add_hyperlink(paragraph, href, link_text)
                else:
                    run = paragraph.add_run(link_text)
                    if color:
                        run.font.color.rgb = color
            elif child.name == "span":
                child_classes = _get_element_classes(child)
                child_color = color
                child_italic = italic
                child_bold = bold
                child_font = font_name
                child_size = font_size
                if "role-date" in child_classes:
                    child_color = COLOR_MUTED_BLUE
                    child_font = "Arial"
                    child_size = Pt(9)
                elif "role-title" in child_classes:
                    child_bold = True
                    child_color = COLOR_DARK_NAVY
                elif "company" in child_classes:
                    child_italic = True
                    child_color = COLOR_DARK_BLUE
                _add_inline_runs(
                    paragraph, child,
                    bold=child_bold, italic=child_italic,
                    font_name=child_font, font_size=child_size,
                    color=child_color,
                )
            else:
                _add_inline_runs(
                    paragraph, child,
                    bold=bold, italic=italic,
                    font_name=font_name, font_size=font_size,
                    color=color,
                )


def _process_element(document, element):
    """Process a single HTML element and add it to the document."""
    if isinstance(element, NavigableString):
        text = str(element).strip()
        if text:
            p = document.add_paragraph()
            run = p.add_run(text)
            run.font.color.rgb = COLOR_BODY
        return

    if not isinstance(element, Tag):
        return

    tag_name = element.name

    if tag_name == "h1":
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        text = element.get_text().strip()
        run = p.add_run(text.upper())
        run.font.name = "Georgia"
        run.font.size = Pt(22)
        run.font.color.rgb = COLOR_DARK_NAVY
        run.bold = False
        p.paragraph_format.space_after = Pt(8)
        _add_bottom_border(p, "0d1b2a", size=8)

    elif tag_name == "h2":
        p = document.add_paragraph()
        text = element.get_text().strip()
        run = p.add_run(text.upper())
        run.font.name = "Georgia"
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_DARK_NAVY
        run.bold = False
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(8)
        _add_bottom_border(p, "778da9", size=4)

    elif tag_name == "h3":
        p = document.add_paragraph()
        text = element.get_text().strip()
        run = p.add_run(text.upper())
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_DARK_BLUE
        run.bold = True
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)

    elif tag_name == "p":
        p = document.add_paragraph()
        classes = _get_element_classes(element)

        if _is_contact(element):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_inline_runs(
                p, element,
                font_name="Arial", font_size=Pt(9), color=COLOR_MUTED_BLUE,
            )
            p.paragraph_format.space_after = Pt(15)
        elif _is_profile(element):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _add_inline_runs(p, element, italic=True, color=COLOR_DARK_BLUE)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.line_spacing = Pt(15)
        elif "company" in classes:
            _add_inline_runs(p, element, italic=True, color=COLOR_DARK_BLUE)
            p.paragraph_format.space_after = Pt(4)
        elif "role-header" in classes:
            _add_inline_runs(p, element, color=COLOR_BODY)
            p.paragraph_format.space_after = Pt(3)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _add_inline_runs(p, element, color=COLOR_BODY)
            p.paragraph_format.space_after = Pt(6)

    elif tag_name in ("ul", "ol"):
        for li in element.find_all("li", recursive=False):
            style = "List Bullet" if tag_name == "ul" else "List Number"
            p = document.add_paragraph(style=style)
            _add_inline_runs(p, li, color=COLOR_BODY)
            p.paragraph_format.space_after = Pt(3)

    elif tag_name == "table":
        rows = element.find_all("tr")
        if not rows:
            return
        max_cols = max(len(row.find_all(["td", "th"])) for row in rows)
        if max_cols == 0:
            return
        table = document.add_table(rows=len(rows), cols=max_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, row in enumerate(rows):
            cells = row.find_all(["td", "th"])
            for j, cell in enumerate(cells):
                if j < max_cols:
                    tc = table.rows[i].cells[j]
                    tc.text = ""
                    p = tc.paragraphs[0]
                    _add_inline_runs(p, cell, color=COLOR_BODY)
                    if j == 0:
                        for run in p.runs:
                            run.bold = True
                            run.font.color.rgb = COLOR_DARK_NAVY

    elif tag_name == "hr":
        p = document.add_paragraph()
        _add_bottom_border(p, "778da9", size=4)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(10)

    elif tag_name == "div":
        classes = _get_element_classes(element)
        if "role-header" in classes:
            # Render role-header as a single paragraph with title left-aligned and date right-aligned
            p = document.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            # Use a right-aligned tab stop at the right margin
            section = document.sections[-1]
            tab_pos = section.page_width - section.left_margin - section.right_margin
            tab_stops = p.paragraph_format.tab_stops
            tab_stops.add_tab_stop(tab_pos, WD_TAB_ALIGNMENT.RIGHT)
            title_span = element.find(class_="role-title")
            date_span = element.find(class_="role-date")
            if title_span:
                run = p.add_run(title_span.get_text())
                run.bold = True
                run.font.color.rgb = COLOR_DARK_NAVY
                run.font.name = "Georgia"
                run.font.size = Pt(10)
            if date_span:
                run = p.add_run("\t")
                run = p.add_run(date_span.get_text())
                run.font.color.rgb = COLOR_MUTED_BLUE
                run.font.name = "Arial"
                run.font.size = Pt(9)
        else:
            for child in element.children:
                _process_element(document, child)

    elif tag_name in ("html", "body", "head", "style", "meta", "title", "link", "script"):
        if tag_name in ("html", "body"):
            for child in element.children:
                _process_element(document, child)
    else:
        # Unknown tags: try to process children
        for child in element.children:
            _process_element(document, child)


@ai_function(name="generate_docx", description="Generate a DOCX file from HTML content")
def generate_docx(
    html_content: Annotated[str, Field(description="The HTML content to convert to DOCX")],
    output_path: Annotated[str, Field(description="The file path where the DOCX should be saved")],
) -> str:
    """
    Generate a DOCX file from HTML content.

    Returns a message indicating success and the output path.
    """
    path = Path(output_path)

    # Ensure the parent directory exists
    path.parent.mkdir(parents=True, exist_ok=True)

    # Create document with A4 page size and matching margins
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    # Set default font
    style = document.styles["Normal"]
    font = style.font
    font.name = "Georgia"
    font.size = Pt(10)
    font.color.rgb = COLOR_BODY

    # Parse HTML and build document
    soup = BeautifulSoup(html_content, "html.parser")

    # Find the main content container, or use body, or use root
    container = soup.find("div", class_="cv-container") or soup.find("body") or soup

    for child in container.children:
        _process_element(document, child)

    document.save(str(path))

    return f"DOCX successfully generated at: {output_path}"
