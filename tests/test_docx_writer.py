"""Tests for the DOCX generation tool."""

from pathlib import Path

from docx import Document

from cv_creator.tools.docx_writer import generate_docx


class TestGenerateDocx:
    def test_generate_docx_creates_file(self, tmp_path):
        output_path = str(tmp_path / "test_cv.docx")
        html = "<html><body><h1>John Doe</h1><p>Software Engineer</p></body></html>"

        result = generate_docx(html, output_path)

        assert Path(output_path).exists()
        assert Path(output_path).stat().st_size > 0
        assert "successfully generated" in result
        assert output_path in result

    def test_generate_docx_creates_parent_dirs(self, tmp_path):
        output_path = str(tmp_path / "nested" / "dir" / "cv.docx")
        html = "<html><body><h1>Jane Doe</h1></body></html>"

        result = generate_docx(html, output_path)

        assert Path(output_path).exists()
        assert "successfully generated" in result

    def test_heading_text_not_concatenated_with_body(self, tmp_path):
        """Verify that heading text is separate from following paragraph text."""
        output_path = str(tmp_path / "test_cv.docx")
        html = """
        <html><body>
            <h1>John Doe</h1>
            <p class="contact">john@example.com | +1234567890</p>
            <h2>Profile</h2>
            <p class="profile">Senior Software Engineer with 10 years experience.</p>
            <h2>Experience</h2>
            <p><strong>Lead Engineer</strong></p>
            <p class="company">EPAM Switzerland</p>
            <p>April 2023 - Present</p>
        </body></html>
        """

        generate_docx(html, output_path)

        doc = Document(output_path)
        texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        # The h1 text should not be concatenated with contact info
        assert any("JOHN DOE" in t and "john@example.com" not in t for t in texts), (
            f"H1 should be separate from contact info. Paragraphs: {texts}"
        )

        # "Profile" heading should be its own paragraph, not merged with body
        assert any(t == "PROFILE" for t in texts), (
            f"H2 'Profile' should be its own paragraph. Paragraphs: {texts}"
        )

        # Company name should not be concatenated with dates
        assert any("EPAM Switzerland" in t and "April 2023" not in t for t in texts), (
            f"Company should be separate from dates. Paragraphs: {texts}"
        )

    def test_lists_are_rendered(self, tmp_path):
        output_path = str(tmp_path / "test_cv.docx")
        html = """
        <html><body>
            <h2>Skills</h2>
            <ul>
                <li>Python</li>
                <li>JavaScript</li>
            </ul>
        </body></html>
        """

        generate_docx(html, output_path)

        doc = Document(output_path)
        texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        assert "Python" in texts
        assert "JavaScript" in texts

    def test_inline_formatting(self, tmp_path):
        output_path = str(tmp_path / "test_cv.docx")
        html = "<html><body><p><strong>Bold</strong> and <em>italic</em> text</p></body></html>"

        generate_docx(html, output_path)

        doc = Document(output_path)
        # Find the paragraph with our content
        for p in doc.paragraphs:
            if "Bold" in p.text:
                runs = p.runs
                bold_runs = [r for r in runs if r.bold and "Bold" in r.text]
                italic_runs = [r for r in runs if r.italic and "italic" in r.text]
                assert bold_runs, "Expected bold formatting on 'Bold'"
                assert italic_runs, "Expected italic formatting on 'italic'"
                break
        else:
            raise AssertionError("Paragraph with 'Bold' text not found")
